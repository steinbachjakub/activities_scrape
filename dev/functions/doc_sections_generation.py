from pathlib import Path
import pandas as pd
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# FILE PATHS
FILE_PATH_ACTIVITIES = Path(".", ".", ".", "data", "activities.csv")

# LOADING SCRAPED DATA
df_activities = pd.read_csv(FILE_PATH_ACTIVITIES)

def generate_doc_intro(document, styles, date_from, date_to):
    title_style, header_style, text_style = styles
    document.add_paragraph('Monthly Activities Report', style=title_style)
    document.add_paragraph(
        f'Dear network, this is the regular report of events submitted in ESN Activities. Below, you can find some highlights from the past month (for the period of {date_from} to {date_to}).',
        style=text_style)
    p = document.add_paragraph(style=text_style)
    p.add_run(
        "Are you curious about how your National Organisation or your local sections fare? Would you like us to create a similar report for you? No problem! Just contact us at ").bold = True
    p.add_run("social-impact@esn.org").italic = True
    p.add_run(".").bold = True
    document.add_paragraph('Few Numbers First', style=header_style)
    document.add_paragraph(
        f"In the past month, total of {df_activities.shape[0]:,} events were organised. The total of {df_activities['participants'].sum():,} participants joined our events which results in the average of {df_activities['participants'].mean():.2f} participants per one event.",
        style=text_style)
    document.add_paragraph(
        "You can see the comparison of event and participant numbers in the maps below.",
        style=text_style)

def generate_doc_section(document, styles, text=None, title=None):
    title_style, header_style, text_style = styles
    if (text is not None) and (title is not None):
        document.add_paragraph(title, style=header_style)
        document.add_paragraph(text, style=text_style)
    document.add_picture("fig_temp.png", height=Cm(8))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.paragraphs[-1].paragraph_format.keep_with_next = False

def generate_doc_outro(document, styles):
    title_style, header_style, text_style = styles
    document.add_paragraph("That's all from us for now. See you next month!", style=text_style)
    document.add_paragraph("Yours faithfully,\nSocial Impact Team", style=text_style)
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT